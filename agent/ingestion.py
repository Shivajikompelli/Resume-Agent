"""
agent/ingestion.py — Resume Ingestion Service

Handles PDF, DOCX, TXT file parsing and text normalization.
Supports OCR fallback for scanned PDFs.
"""

import re
import logging
import unicodedata
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)


class IngestionService:
    """
    Extracts and normalises plain text from resume files.
    Supported formats: PDF, DOCX, TXT
    """

    MAX_FILE_SIZE_MB = 10

    def extract(self, file_path: str | Path) -> Tuple[str, dict]:
        """
        Extract text from a resume file.

        Returns:
            text: Cleaned, normalised text content
            meta: Metadata dict (format, ocr_used, char_count, etc.)
        """
        path = Path(file_path)
        size_mb = path.stat().st_size / (1024 * 1024)

        if size_mb > self.MAX_FILE_SIZE_MB:
            raise ValueError(f"File too large: {size_mb:.1f} MB (max {self.MAX_FILE_SIZE_MB} MB)")

        suffix = path.suffix.lower()
        meta = {"format": suffix, "ocr_used": False, "file_size_mb": round(size_mb, 2)}

        if suffix == ".pdf":
            text, ocr_used = self._extract_pdf(path)
            meta["ocr_used"] = ocr_used
        elif suffix in (".docx", ".doc"):
            text = self._extract_docx(path)
        elif suffix == ".txt":
            text = path.read_text(encoding="utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported file type: {suffix}. Use PDF, DOCX, or TXT.")

        text = self._clean(text)
        meta["char_count"] = len(text)
        meta["word_count"] = len(text.split())

        if len(text.strip()) < 50:
            raise ValueError("Could not extract meaningful text from the file. "
                             "Please check the file is not corrupted or password-protected.")

        return text, meta

    # ── Private helpers ──────────────────────────────────────────

    def _extract_pdf(self, path: Path) -> Tuple[str, bool]:
        """Extract text from PDF; fall back to OCR for scanned pages."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")

        pages_text = []
        ocr_used = False

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if len(text.strip()) < 30:
                    # Empty or near-empty page — try OCR
                    ocr_text = self._ocr_page(page)
                    if ocr_text:
                        pages_text.append(ocr_text)
                        ocr_used = True
                    else:
                        pages_text.append(text)
                else:
                    pages_text.append(text)

        return "\n".join(pages_text), ocr_used

    def _ocr_page(self, page) -> str:
        """Attempt OCR on a PDF page using Tesseract."""
        try:
            import pytesseract
            from PIL import Image
            img = page.to_image(resolution=300).original
            return pytesseract.image_to_string(img, lang="eng")
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return ""

    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX preserving reading order."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document(str(path))
        parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    def _clean(self, text: str) -> str:
        """Normalise and clean extracted text."""
        # Normalise unicode (fix mojibake, ligatures, etc.)
        text = unicodedata.normalize("NFKC", text)

        # Replace tabs with spaces
        text = text.replace("\t", " ")

        # Collapse multiple blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Strip trailing whitespace from lines
        lines = [line.rstrip() for line in text.split("\n")]
        text = "\n".join(lines)

        # Remove non-printable control characters (keep \n)
        text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\x80-\xFF]", "", text)

        return text.strip()
